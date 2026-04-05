import os
import json
import pdfplumber
import docx
import markdownify
import logging
from flask import current_app
from datetime import datetime
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage, SystemMessage

# OCR
import pytesseract
from PIL import Image

logger = logging.getLogger(__name__)

pytesseract.pytesseract.tesseract_cmd = r"C:/Program Files/Tesseract-OCR/tesseract.exe"


# ---------------------------------------------------------
# CLEAN HELPERS
# ---------------------------------------------------------
def clean_utf(text):
    if not isinstance(text, str):
        return text
    # Fix common mojibake (UTF-8 bytes decoded as latin-1) WITHOUT dropping accents.
    # We only attempt the conversion when we detect typical artifacts.
    try:
        if any(bad in text for bad in ["Ã", "Â", "â€", "â€™", "â€œ", "â€\x9d", "Å", "ðŸ", "�"]):
            try:
                fixed = text.encode("latin1").decode("utf-8")
                return fixed
            except Exception:
                pass
        return text
    except Exception:
        return text


def clean_json(x):
    if isinstance(x, dict):
        return {k: clean_json(v) for k, v in x.items()}
    if isinstance(x, list):
        return [clean_json(v) for v in x]
    if isinstance(x, str):
        return clean_utf(x)
    return x


# ---------------------------------------------------------
# AAP DETECTOR WITH OCR (WORKS WITH ESB PDF)
# ---------------------------------------------------------
import base64
import pdfplumber
import io
import io
from flask import current_app

def _find_aap_header(words, page_num):
    """Find AAP header with multiple strategies."""
    # Strategy 1: Exact "AAP" match
    aap_words = [w for w in words if w["text"].startswith("AAP")]
    if aap_words:
        logger.info(f"Found AAP header on page {page_num + 1}")
        return aap_words[0]

    # Strategy 2: Case-insensitive
    aap_words = [w for w in words if w["text"].upper().startswith("AAP")]
    if aap_words:
        logger.info(f"Found AAP header (case-insensitive) on page {page_num + 1}")
        return aap_words[0]

    return None


def extract_aap_from_pdf(pdf_path):
    """
    Automatically detects the AAP table by finding the text 'AAP' and cropping around it.
    Then uses Gemini Vision to detect X marks. PDF only.
    """
    # DOCX files cannot be processed for AAP (requires visual PDF analysis)
    if not pdf_path.lower().endswith('.pdf'):
        logger.info(f"AAP extraction skipped for non-PDF file: {pdf_path}")
        return []

    try:
        with pdfplumber.open(pdf_path) as pdf:

            if len(pdf.pages) < 1:
                logger.warning("PDF has no pages")
                return []

            # Try multiple pages: page 3 first (index 2), then 2, 4, 5, and finally all pages
            pages_to_check = [2, 1, 3, 4]

            # Add remaining pages if PDF is longer
            for i in range(len(pdf.pages)):
                if i not in pages_to_check:
                    pages_to_check.append(i)

            aap = None
            page = None
            found_page_idx = -1

            for page_idx in pages_to_check:
                if page_idx >= len(pdf.pages):
                    continue

                page = pdf.pages[page_idx]
                words = page.extract_words()
                aap = _find_aap_header(words, page_idx)

                if aap:
                    found_page_idx = page_idx
                    break

            if not aap:
                logger.warning(f"No AAP header found on any page (checked {min(len(pages_to_check), len(pdf.pages))} pages)")
                return []

            # Take the first AAP word position
            logger.info(f"Processing AAP table from page {found_page_idx + 1}")

            x0 = aap["x0"] - 20     # expand a bit left
            top = aap["top"] - 40   # go higher
            x1 = page.width - 20
            bottom = aap["top"] + 180  # go lower

            bbox = (x0, top, x1, bottom)
            cropped = page.crop(bbox)

            # Save debug crop
            img = cropped.to_image(resolution=250).original
            img.save("debug_aap_crop.png")

            # Convert to base64
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            b64_img = base64.b64encode(buf.getvalue()).decode()

        # ---- CALL GEMINI VISION MODEL ----
        api_key = current_app.config.get("GOOGLE_API_KEY")
        if not api_key:
             raise ValueError("Google API key not configured")
             
        model = current_app.config.get("GEMINI_MODEL", "gemini-2.5-flash")
        
        llm = ChatGoogleGenerativeAI(
            model=model,
            google_api_key=api_key,
            temperature=0
        )

        prompt_str = """
You will see a table row with AAP numbers (1–12). 
Some boxes contain an X. Identify only which AAP numbers contain an X.

Respond ONLY as JSON inside <json></json>:

<json>
[
  {"AAP#": 7, "selected": true},
  {"AAP#": 8, "selected": true}
]
</json>
"""
        messages = [
            HumanMessage(content=[
                {"type": "text", "text": prompt_str},
                {
                    "type": "image_url",
                    "image_url": f"data:image/png;base64,{b64_img}"
                }
            ])
        ]

        response = llm.invoke(messages)

        content = response.content

        # Use robust JSON extraction
        try:
            # Try to extract JSON using the same multi-strategy approach
            json_data = _extract_json_from_llm_response(content)
            if json_data and isinstance(json_data, list):
                return json_data
            elif json_data and isinstance(json_data, dict):
                # If it's a dict, try to extract an array from it
                for key in json_data:
                    if isinstance(json_data[key], list):
                        return json_data[key]
            return []
        except Exception as e:
            logger.error(f"Failed to parse AAP vision response: {e}")
            return []

    except Exception as e:
        logger.error(f"AAP Vision LLM error: {e}")
        return []

import re


def _clean_json_string(json_str: str) -> str:
    """
    Clean common LLM JSON formatting issues.
    Removes markdown, XML tags, fixes double braces, removes trailing commas.
    """
    import re

    # Remove markdown code blocks
    json_str = re.sub(r'```(?:json)?', '', json_str)

    # Remove XML-style tags if present
    json_str = re.sub(r'</?json>', '', json_str)

    # Fix double braces (Gemini sometimes escapes braces)
    json_str = json_str.replace("{{", "{").replace("}}", "}")

    # Remove trailing commas before closing brackets
    json_str = re.sub(r',\s*([}\]])', r'\1', json_str)

    # Clean up whitespace
    json_str = json_str.strip()

    return json_str


def _extract_json_from_llm_response(content: str) -> dict:
    """
    Module-level helper to extract JSON from LLM response with multiple fallback strategies.
    Used by both main extraction and AAP vision detection.
    Reordered to handle mixed format responses (markdown + XML) correctly.
    """
    import re

    # Log preview of raw response for debugging
    logger.debug(f"Raw LLM response preview: {content[:200]}...")
    logger.debug(f"Response contains <json> tag: {'<json>' in content}")
    logger.debug(f"Response contains markdown block: {'```' in content}")

    # Strategy 1a: Regex for JSON objects
    logger.debug("Trying Strategy 1a: Regex for JSON objects")
    try:
        pattern = r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}'
        match = re.search(pattern, content, re.DOTALL)
        if match:
            json_str = _clean_json_string(match.group(0))
            result = json.loads(json_str)
            logger.debug("Strategy 1a (object regex) succeeded")
            return result
    except Exception as e:
        logger.debug(f"Strategy 1a failed: {e}")

    # Strategy 1b: Regex for JSON arrays with objects
    logger.debug("Trying Strategy 1b: Regex for JSON arrays with objects")
    try:
        pattern = r'\[\s*\{[\s\S]*?\}\s*\]'
        match = re.search(pattern, content, re.DOTALL)
        if match:
            json_str = _clean_json_string(match.group(0))
            result = json.loads(json_str)
            logger.debug("Strategy 1b (array regex) succeeded")
            return result
    except Exception as e:
        logger.debug(f"Strategy 1b failed: {e}")

    # Strategy 1c: Regex for simple arrays
    logger.debug("Trying Strategy 1c: Regex for simple arrays")
    try:
        pattern = r'\[[\s\S]*?\]'
        match = re.search(pattern, content, re.DOTALL)
        if match:
            json_str = _clean_json_string(match.group(0))
            result = json.loads(json_str)
            logger.debug("Strategy 1c (simple array regex) succeeded")
            return result
    except Exception as e:
        logger.debug(f"Strategy 1c failed: {e}")

    # Strategy 2: Markdown code blocks
    logger.debug("Trying Strategy 2: Markdown code blocks")
    code_block_match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", content)
    if code_block_match:
        try:
            json_str = _clean_json_string(code_block_match.group(1))
            result = json.loads(json_str)
            logger.debug("Strategy 2 (markdown code block) succeeded")
            return result
        except Exception as e:
            logger.debug(f"Strategy 2 failed: {e}")

    # Strategy 3: XML tags WITH cleanup (FIX FOR THE BUG)
    logger.debug("Trying Strategy 3: XML tags with cleanup")
    if "<json>" in content:
        try:
            json_str = content.split("<json>")[1].split("</json>")[0].strip()
            # CRITICAL: Apply cleanup to remove embedded markdown
            json_str = _clean_json_string(json_str)
            result = json.loads(json_str)
            logger.debug("Strategy 3 (XML tags with cleanup) succeeded")
            return result
        except Exception as e:
            logger.debug(f"Strategy 3 failed: {e}")

    # Strategy 4: Raw JSON (starts with { or [)
    logger.debug("Trying Strategy 4: Raw JSON")
    stripped = content.strip()
    if stripped.startswith("{") or stripped.startswith("["):
        try:
            json_str = _clean_json_string(stripped)
            result = json.loads(json_str)
            logger.debug("Strategy 4 (raw JSON) succeeded")
            return result
        except Exception as e:
            logger.debug(f"Strategy 4 failed: {e}")

    # Strategy 5a: Outermost braces
    logger.debug("Trying Strategy 5a: Outermost braces")
    first_brace = content.find("{")
    last_brace = content.rfind("}")
    if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
        try:
            json_str = _clean_json_string(content[first_brace:last_brace + 1])
            result = json.loads(json_str)
            logger.debug("Strategy 5a (outermost braces) succeeded")
            return result
        except Exception as e:
            logger.debug(f"Strategy 5a failed: {e}")

    # Strategy 5b: Outermost brackets (FIX FOR ARRAY SUPPORT)
    logger.debug("Trying Strategy 5b: Outermost brackets")
    first_bracket = content.find("[")
    last_bracket = content.rfind("]")
    if first_bracket != -1 and last_bracket != -1 and last_bracket > first_bracket:
        try:
            json_str = _clean_json_string(content[first_bracket:last_bracket + 1])
            result = json.loads(json_str)
            logger.debug("Strategy 5b (outermost brackets) succeeded")
            return result
        except Exception as e:
            logger.debug(f"Strategy 5b failed: {e}")

    logger.error("LLM returned no parseable JSON (tried all strategies)")
    logger.error(f"Full content: {content[:1000]}...")
    return {}


def _extract_json_block(tagged_text: str, tag: str = "json") -> str:
    """
    Robust extraction of JSON content from LLM response.
    Returns cleaned JSON string (not parsed object).
    Reordered strategies to handle mixed format responses correctly.
    """
    if not tagged_text:
        raise ValueError("Empty LLM response")

    text = tagged_text.strip()
    logger.debug(f"Extracting JSON block preview: {text[:200]}...")

    import re

    # Strategy 1a: Regex for JSON objects
    logger.debug("Trying Strategy 1a: Regex for JSON objects")
    try:
        pattern = r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}'
        match = re.search(pattern, text, re.DOTALL)
        if match:
            json_str = _clean_json_string(match.group(0))
            logger.debug("Strategy 1a (object regex) succeeded")
            return json_str
    except Exception as e:
        logger.debug(f"Strategy 1a failed: {e}")

    # Strategy 1b: Regex for JSON arrays with objects
    logger.debug("Trying Strategy 1b: Regex for JSON arrays with objects")
    try:
        pattern = r'\[\s*\{[\s\S]*?\}\s*\]'
        match = re.search(pattern, text, re.DOTALL)
        if match:
            json_str = _clean_json_string(match.group(0))
            logger.debug("Strategy 1b (array regex) succeeded")
            return json_str
    except Exception as e:
        logger.debug(f"Strategy 1b failed: {e}")

    # Strategy 1c: Regex for simple arrays
    logger.debug("Trying Strategy 1c: Regex for simple arrays")
    try:
        pattern = r'\[[\s\S]*?\]'
        match = re.search(pattern, text, re.DOTALL)
        if match:
            json_str = _clean_json_string(match.group(0))
            logger.debug("Strategy 1c (simple array regex) succeeded")
            return json_str
    except Exception as e:
        logger.debug(f"Strategy 1c failed: {e}")

    # Strategy 2: Markdown code blocks
    logger.debug("Trying Strategy 2: Markdown code blocks")
    code_block_pattern = r"```(?:json)?\s*([\s\S]*?)\s*```"
    match = re.search(code_block_pattern, text)
    if match:
        json_str = _clean_json_string(match.group(1))
        logger.debug("Strategy 2 (markdown code block) succeeded")
        return json_str

    # Strategy 3: XML tags WITH cleanup
    logger.debug("Trying Strategy 3: XML tags with cleanup")
    start_tag = f"<{tag}>"
    end_tag = f"</{tag}>"
    s_idx = text.find(start_tag)
    e_idx = text.find(end_tag)

    if s_idx != -1 and e_idx != -1 and e_idx > s_idx:
        json_str = text[s_idx + len(start_tag):e_idx].strip()
        # CRITICAL: Apply cleanup to remove embedded markdown
        json_str = _clean_json_string(json_str)
        logger.debug("Strategy 3 (XML tags with cleanup) succeeded")
        return json_str

    # Strategy 4: Raw JSON (starts with { or [)
    logger.debug("Trying Strategy 4: Raw JSON")
    if (text.startswith("{") and text.endswith("}")) or (text.startswith("[") and text.endswith("]")):
        json_str = _clean_json_string(text)
        logger.debug("Strategy 4 (raw JSON) succeeded")
        return json_str

    # Strategy 5a: Outermost braces
    logger.debug("Trying Strategy 5a: Outermost braces")
    s_brace = text.find("{")
    e_brace = text.rfind("}")
    if s_brace != -1 and e_brace != -1 and e_brace > s_brace:
        json_str = _clean_json_string(text[s_brace:e_brace + 1])
        logger.debug("Strategy 5a (outermost braces) succeeded")
        return json_str

    # Strategy 5b: Outermost brackets
    logger.debug("Trying Strategy 5b: Outermost brackets")
    s_bracket = text.find("[")
    e_bracket = text.rfind("]")
    if s_bracket != -1 and e_bracket != -1 and e_bracket > s_bracket:
        json_str = _clean_json_string(text[s_bracket:e_bracket + 1])
        logger.debug("Strategy 5b (outermost brackets) succeeded")
        return json_str

    # If all fail, raise error but log the content for debugging
    logger.warning(f"Failed to extract JSON. Content start: {text[:100]}...")
    raise ValueError("No valid JSON block found in response")


def _safe_json_loads(s: str):
    try:
        return json.loads(s)
    except Exception as e:
        raise ValueError(f"JSON parse failed: {e}")


# ---------------------------------------------------------
# MAIN EXTRACTOR
# ---------------------------------------------------------
class SyllabusTNService:

    @staticmethod
    def _client(temperature=0):
        api = current_app.config.get("GOOGLE_API_KEY")
        model = current_app.config.get("GEMINI_MODEL", "gemini-2.5-flash")
        if not api:
            raise ValueError("Google API key missing")
             
        return ChatGoogleGenerativeAI(
            model=model,
            google_api_key=api,
            temperature=temperature
        )

    # -----------------------
    # FILE PARSER (PDF & DOCX)
    # -----------------------
    @staticmethod
    def extract_text(path):
        ext = path.lower().split(".")[-1]

        # PDF
        if ext == "pdf":
            with pdfplumber.open(path) as pdf:
                pages = []
                for p in pdf.pages:
                    t = p.extract_text()
                    if t:
                        pages.append(markdownify.markdownify(t, heading_style="ATX"))
            return "\n".join(pages)

        # DOCX
        if ext in ["doc", "docx"]:
            d = docx.Document(path)
            parts = []
            # Paragraphs (may be sparse in table-heavy TN fiches)
            para_text = "\n".join([clean_utf(p.text) for p in d.paragraphs if p.text.strip()])
            if para_text:
                parts.append(para_text)
            # Tables — TN fiches store all content in a single merged table.
            # Use XML identity (element pointer) to deduplicate merged cells.
            for tbl in d.tables:
                seen_ptrs = set()
                for row in tbl.rows:
                    row_parts = []
                    for cell in row.cells:
                        ptr = cell._tc  # lxml element — same object for merged cells
                        if ptr in seen_ptrs:
                            continue
                        seen_ptrs.add(ptr)
                        ct = clean_utf(cell.text.strip())
                        if ct:
                            row_parts.append(ct)
                    if row_parts:
                        parts.append(" | ".join(row_parts))
            return "\n".join(parts)

        return ""

    # -----------------------
    # ONE SINGLE LLM CALL
    # -----------------------
    @staticmethod
    def ask_llm(text, retry_count=0, max_retries=2):

        llm = SyllabusTNService._client(temperature=0)

        prompt_str = f"""
You are a strict JSON extractor for Tunisian university syllabus.

Extract ALL the following IN ONE SINGLE JSON:

<json>
{{
  "administrative": {{
    "module_name": "",
    "code_ue": "",
    "code_ecue": "",
    "field": "",
    "department": "",
    "option": "",
    "volume_presentiel": "",
    "volume_personnel": "",
    "coefficient": 0,
    "credits": 0,
    "responsible": "",
    "teachers": []
  }},
  "aaa": [
    {{"AA#": 1, "description": ""}}
  ],
  "chapters": [
    {{
      "chapter": "",
      "sections": ["", ""]
    }}
  ],
  "evaluation": {{
    "methods": [],
    "criteria": [],
    "measures": [],
    "final_grade_formula": ""
  }},
  "bibliography": []
}}
</json>

Rules:
- Return ONLY JSON inside <json></json>.
- DO NOT invent AAP. AAP is handled separately.
- Respect accents (é, à, è).

TEXT:
{text}
"""
        messages = [
            SystemMessage(content="Return ONLY JSON inside <json></json>."),
            HumanMessage(content=prompt_str)
        ]

        try:
            response = llm.invoke(messages)
            content = response.content

            # Use robust JSON extraction
            data = _extract_json_from_llm_response(content)

            # Check if extraction returned empty dict and retry if needed
            if not data and retry_count < max_retries:
                logger.warning(f"Empty extraction result, retrying ({retry_count + 1}/{max_retries})...")
                import time
                time.sleep(30)  # Wait 30 seconds before retry
                return SyllabusTNService.ask_llm(text, retry_count + 1, max_retries)

            if not data:
                logger.error("LLM returned empty content after all retries")
                return {}

            return clean_json(data)

        except Exception as e:
            logger.error(f"LLM invocation failed: {e}")
            if retry_count < max_retries:
                logger.warning(f"Retrying due to error ({retry_count + 1}/{max_retries})...")
                import time
                time.sleep(30)
                return SyllabusTNService.ask_llm(text, retry_count + 1, max_retries)
            return {}

    # -----------------------
    # MASTER EXTRACTION
    # -----------------------
    @staticmethod
    def extract_tn_syllabus(file_path):

        text = SyllabusTNService.extract_text(file_path)

        # ONE LLM CALL for everything except AAP
        llm_data = SyllabusTNService.ask_llm(text)

        # AAP from OCR detector
        aap = extract_aap_from_pdf(file_path)

        # Build final JSON
        result = {
            "administrative": llm_data.get("administrative", {}),
            "aaa": llm_data.get("aaa", []),
            "aap": aap,
            "mapping": [],
            "chapters": llm_data.get("chapters", []),
            "evaluation": llm_data.get("evaluation", {}),
            "bibliography": llm_data.get("bibliography", []),
        }

        return clean_json(result)

    @staticmethod
    def classify_chapters_sections_to_aaa(extracted: dict) -> dict:
        """
        Input: extracted TN syllabus JSON (administrative/aaa/chapters/...)
        Output: extracted + new field extracted["aaa_classification"]
        """
        llm = SyllabusTNService._client(temperature=0)

        # Build context strings from input 'extracted'
        aaa_list = extracted.get("aaa", [])
        if not aaa_list:
            return extracted  # nothing to classify against

        aaa_text = "\n".join([f"AA {entry.get('AA#', '?')}: {entry.get('description', '')}"
                              for entry in aaa_list])

        chapters_list = extracted.get("chapters", [])
        if not chapters_list:
            return extracted # nothing to classify

        chapters_text = json.dumps(chapters_list, indent=2, ensure_ascii=False)

        prompt_str = f"""
[INST]
Tu es un assistant intelligent qui fait une classification sémantique (pas juste mots-clés).
Objectif: associer des sections et chapitres à des AAA (Acquis d’apprentissage) pertinents.

Voici la liste des AA disponibles:
{aaa_text}

Voici le plan du cours (chapitres + sections) à classer:
{chapters_text}

Tâche:
1) Pour chaque CHAPTER#, donne la liste des AA# correspondants (au moins 1).
2) Pour chaque SECTION# (i.j), donne aussi la liste des AA# correspondants (au moins 1).
3) Utilise le sens global (concepts implicites) pas seulement les mots.
4) Si un item correspond à plusieurs AA, retourne-les tous.

Règles de sortie:
- Retourne uniquement un JSON valide dans <json></json>
- Format EXACT:

<json>
{{
  "chapters": [
    {{
      "chapter_index": 1,
      "chapter_title": "...",
      "AA#": [1,2],
      "AADescription": ["...", "..."],
      "sections": [
        {{
          "section_index": "1.1",
          "section_title": "...",
          "AA#": [2],
          "AADescription": ["..."]
        }}
      ]
    }}
  ]
}}
</json>
[/INST]
"""
        messages = [HumanMessage(content=prompt_str)]

        response = llm.invoke(messages)

        content = response.content
        block = _extract_json_block(content, "json")
        classification = _safe_json_loads(block)

        # attach to extracted (non destructive)
        extracted = dict(extracted)
        extracted["aaa_classification"] = classification
        return clean_json(extracted)