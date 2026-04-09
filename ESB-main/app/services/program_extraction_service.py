"""
Service for extracting AAP, Competences and their relationships
from a Word (.docx) "fiche descriptive de formation" document.
"""
import logging
import re
from typing import Dict, List, Optional, Tuple

import docx

logger = logging.getLogger(__name__)

# ── Regex patterns ───────────────────────────────────────────────────────────
_RE_AAP_CODE = re.compile(r'AAP\s*(\d+)', re.IGNORECASE)
_RE_COMP_CODE = re.compile(r'C\s*(\d+)', re.IGNORECASE)
_SKIP_KEYWORDS = ('instructions', 'indiquer', 'identifier')


# ── Helpers ──────────────────────────────────────────────────────────────────

def _cell_text(cell) -> str:
    """Return stripped text of a table cell, collapsing whitespace."""
    return ' '.join(cell.text.split()).strip()


def _row_texts(row) -> List[str]:
    """Return list of stripped cell texts for a row."""
    return [_cell_text(c) for c in row.cells]


def _should_skip_row(row_texts: List[str]) -> bool:
    """Return True if the row is an instruction/header row to ignore."""
    first = row_texts[0].lower() if row_texts else ''
    return any(kw in first for kw in _SKIP_KEYWORDS)


def _find_tables(doc: docx.Document) -> Tuple[Optional[object], Optional[object], Optional[object]]:
    """
    Detect the competences table, AAP table, and matrix table by content
    rather than hardcoded indices.

    Returns (competences_table, aap_table, matrix_table).
    """
    comp_table = None
    aap_table = None
    matrix_table = None

    for idx, table in enumerate(doc.tables):
        rows = table.rows
        if len(rows) < 2:
            continue

        # Collect text from the first few rows for detection
        sample_texts = []
        for r in rows[:min(5, len(rows))]:
            sample_texts.extend([_cell_text(c).lower() for c in r.cells])
        joined = ' '.join(sample_texts)

        # ── Matrix table: has both "aap" references and "c " competence
        #    references, and typically has many columns (>5)
        num_cols = len(rows[0].cells)
        if num_cols > 5 and 'aap' in joined and re.search(r'\bc\s*\d', joined):
            if matrix_table is None:
                matrix_table = table
                logger.info("Detected matrix table at index %d (%d rows x %d cols)",
                            idx, len(rows), num_cols)
            continue

        # ── AAP table: first-column cells match "AAP <digit>"
        if _RE_AAP_CODE.search(joined) and num_cols <= 5:
            has_aap_rows = False
            for r in rows[2:min(6, len(rows))]:
                rt = _row_texts(r)
                if any(_RE_AAP_CODE.match(t) for t in rt):
                    has_aap_rows = True
                    break
            if has_aap_rows and aap_table is None:
                aap_table = table
                logger.info("Detected AAP table at index %d (%d rows x %d cols)",
                            idx, len(rows), num_cols)
                continue

        # ── Competences table: cells match "C <digit>"
        if _RE_COMP_CODE.search(joined) and num_cols <= 5:
            has_comp_rows = False
            for r in rows[2:min(6, len(rows))]:
                rt = _row_texts(r)
                if any(_RE_COMP_CODE.match(t) for t in rt):
                    has_comp_rows = True
                    break
            if has_comp_rows and comp_table is None:
                comp_table = table
                logger.info("Detected competences table at index %d (%d rows x %d cols)",
                            idx, len(rows), num_cols)
                continue

    return comp_table, aap_table, matrix_table


# ── Extractors ───────────────────────────────────────────────────────────────

def _extract_competences(table) -> List[dict]:
    """Extract competence rows from detected competences table."""
    competences: List[dict] = []
    for row in table.rows:
        texts = _row_texts(row)
        if _should_skip_row(texts):
            continue

        # Find the cell that contains the competence code
        code = None
        code_idx = -1
        for i, t in enumerate(texts):
            m = _RE_COMP_CODE.match(t)
            if m:
                code = f"C {m.group(1)}"
                code_idx = i
                break

        if code is None:
            continue

        # Remaining cells provide description and optionally nature
        remaining = [t for j, t in enumerate(texts) if j != code_idx and t]
        description = remaining[0] if remaining else ''
        nature = remaining[1] if len(remaining) > 1 else ''

        competences.append({
            'code': code,
            'description': description.strip(),
            'nature': nature.strip(),
        })

    logger.info("Extracted %d competences", len(competences))
    return competences


def _extract_aaps(table) -> List[dict]:
    """Extract AAP rows from detected AAP table."""
    aaps: List[dict] = []
    for row in table.rows:
        texts = _row_texts(row)
        if _should_skip_row(texts):
            continue

        code = None
        code_idx = -1
        for i, t in enumerate(texts):
            m = _RE_AAP_CODE.match(t)
            if m:
                code = f"AAP {m.group(1)}"
                code_idx = i
                break

        if code is None:
            continue

        remaining = [t for j, t in enumerate(texts) if j != code_idx and t]
        denomination = remaining[0] if remaining else ''
        description = remaining[1] if len(remaining) > 1 else ''

        aaps.append({
            'code': code,
            'denomination': denomination.strip(),
            'description': description.strip(),
            'order': int(_RE_AAP_CODE.match(code).group(1)),
        })

    logger.info("Extracted %d AAPs", len(aaps))
    return aaps


def _extract_matrix(table) -> List[dict]:
    """
    Extract the competence ↔ AAP relationship matrix.

    Returns a list of dicts:
      [{'competence_code': 'C 1', 'aap_codes': ['AAP 1', 'AAP 3']}, ...]
    """
    rows = table.rows
    if len(rows) < 3:
        logger.warning("Matrix table has fewer than 3 rows; cannot parse headers")
        return []

    # Find the header row that contains AAP column labels.
    # It is typically the last row before data rows start (row with "C " in col 0).
    aap_header_row_idx = None
    aap_columns: Dict[int, str] = {}  # col_index -> AAP code

    for r_idx in range(min(5, len(rows))):
        texts = _row_texts(rows[r_idx])
        found_aap = False
        for c_idx, t in enumerate(texts):
            m = _RE_AAP_CODE.search(t)
            if m:
                found_aap = True
                aap_columns[c_idx] = f"AAP {m.group(1)}"
        if found_aap:
            aap_header_row_idx = r_idx
            break

    if not aap_columns:
        logger.warning("Could not find AAP column headers in matrix table")
        return []

    logger.info("Matrix AAP columns detected: %s", aap_columns)

    matrix: List[dict] = []
    for row in rows[aap_header_row_idx + 1:]:
        texts = _row_texts(row)
        if _should_skip_row(texts):
            continue

        # First cell should contain the competence code
        first_cell = texts[0] if texts else ''
        m = _RE_COMP_CODE.search(first_cell)
        if not m:
            continue
        comp_code = f"C {m.group(1)}"

        # Check which AAP columns are marked with "X"
        linked_aaps = []
        for c_idx, aap_code in aap_columns.items():
            if c_idx < len(texts) and texts[c_idx].strip().upper() == 'X':
                linked_aaps.append(aap_code)

        matrix.append({
            'competence_code': comp_code,
            'aap_codes': linked_aaps,
        })

    logger.info("Extracted matrix with %d competence rows", len(matrix))
    return matrix


def _extract_study_plan(doc) -> List[dict]:
    """
    Extract ECUE (module) names from the study plan tables.
    Returns a list of dicts: [{'name': '...', 'code': '...', 'semester': N, 'ue': '...'}, ...]
    """
    ecues: List[dict] = []
    _RE_SEMESTER = re.compile(r'semestre\s*(\d+)', re.IGNORECASE)
    seen_names = set()
    semester_counter = 0

    for table in doc.tables:
        rows = table.rows
        if len(rows) < 3:
            continue

        # Check row 0 and row 1 headers (some tables have merged header rows)
        header0 = [_cell_text(c).lower() for c in rows[0].cells]
        header1 = [_cell_text(c).lower() for c in rows[1].cells] if len(rows) > 1 else header0
        joined_header = ' '.join(header0 + header1)

        # Must have both "ecue" and "ue" in headers, typical of study plan tables
        if not ('ecue' in joined_header and ('crédit' in joined_header or 'coef' in joined_header)):
            continue

        # Auto-increment semester for each study plan table found
        semester_counter += 1

        # Find the right columns using row 1 (sub-headers)
        ecue_name_col = -1
        ecue_code_col = -1
        ue_name_col = -1

        for i, h in enumerate(header0):
            h_clean = h.strip()
            # "Élément constitutif d'UE (ECUE)" = the name column
            if 'constitutif' in h_clean or (h_clean.startswith('el') and 'ecue' in h_clean):
                ecue_name_col = i
            # "Code de l'ECUE" = code column
            elif 'code' in h_clean and 'ecue' in h_clean:
                ecue_code_col = i
            # "Unité d'enseignement (UE)" = UE name
            elif 'enseignement' in h_clean or (h_clean.startswith('unit') and 'ue' in h_clean):
                ue_name_col = i

        if ecue_name_col == -1:
            continue

        current_ue = ''
        for row in rows[2:]:  # Skip header rows
            texts = _row_texts(row)
            if _should_skip_row(texts):
                continue
            # Skip total/sub-total rows
            first = texts[0].lower() if texts else ''
            if any(kw in first for kw in ('total', 'sous-total', 'sous total', 'n°')):
                continue

            # Update current UE
            if ue_name_col >= 0 and ue_name_col < len(texts) and texts[ue_name_col].strip():
                current_ue = texts[ue_name_col].strip()

            # Get ECUE name
            if ecue_name_col < len(texts):
                ecue_name = texts[ecue_name_col].strip()
                ecue_code = texts[ecue_code_col].strip() if ecue_code_col >= 0 and ecue_code_col < len(texts) else ''
                if ecue_name and ecue_name.lower() not in seen_names:
                    seen_names.add(ecue_name.lower())
                    ecues.append({
                        'name': ecue_name,
                        'code': ecue_code,
                        'semester': semester_counter,
                        'ue': current_ue,
                    })

    logger.info("Extracted %d ECUE (modules) from study plan", len(ecues))
    return ecues


# ── Public API ───────────────────────────────────────────────────────────────

def extract_program_descriptor(file_path: str, program_id: int) -> dict:
    """
    Parse a .docx formation descriptor and extract AAP, Competences,
    and their relationship matrix.

    Args:
        file_path: Path to the .docx file.
        program_id: The program (formation) id for context.

    Returns:
        {
            'aaps': [{'code': 'AAP 1', 'denomination': '...', 'description': '...', 'order': 1}, ...],
            'competences': [{'code': 'C 1', 'description': '...', 'nature': '...'}, ...],
            'matrix': [{'competence_code': 'C 1', 'aap_codes': ['AAP 1', 'AAP 3']}, ...],
        }
    """
    logger.info("Opening document: %s (program_id=%d)", file_path, program_id)
    doc = docx.Document(file_path)
    logger.info("Document loaded — %d tables found", len(doc.tables))

    comp_table, aap_table, matrix_table = _find_tables(doc)

    competences = _extract_competences(comp_table) if comp_table else []
    aaps = _extract_aaps(aap_table) if aap_table else []
    matrix = _extract_matrix(matrix_table) if matrix_table else []
    study_plan = _extract_study_plan(doc)

    if not comp_table:
        logger.warning("Competences table not detected")
    if not aap_table:
        logger.warning("AAP table not detected")
    if not matrix_table:
        logger.warning("Matrix table not detected")

    return {
        'aaps': aaps,
        'competences': competences,
        'matrix': matrix,
        'study_plan': study_plan,
        'teachers': _extract_teachers(doc),
    }


def _extract_teachers(doc) -> List[dict]:
    """
    Extract teacher info from the teacher tables in the document.
    Tables have headers containing 'Nom et Prénom' and 'ECUE'.
    Returns: [{'name': '...', 'grade': '...', 'specialty': '...', 'ecues': ['...']}, ...]
    """
    teachers: List[dict] = []
    seen_names = set()

    for table in doc.tables:
        rows = table.rows
        if len(rows) < 2:
            continue

        header = [_cell_text(c).lower() for c in rows[0].cells]
        joined = ' '.join(header)

        if 'nom' not in joined or 'ecue' not in joined:
            continue
        if 'grade' not in joined and 'dipl' not in joined:
            continue

        # Find column indices
        name_col = -1
        grade_col = -1
        specialty_col = -1
        ecue_col = -1
        establishment_col = -1

        for i, h in enumerate(header):
            if 'nom' in h and 'pr' in h:
                name_col = i
            elif 'grade' in h or 'dipl' in h:
                grade_col = i
            elif 'sp' in h and 'cialit' in h:
                specialty_col = i
            elif 'ecue' in h:
                ecue_col = i
            elif 'etablissement' in h or 'organisation' in h:
                establishment_col = i

        if name_col == -1 or ecue_col == -1:
            continue

        for row in rows[1:]:
            texts = _row_texts(row)
            name = texts[name_col].strip() if name_col < len(texts) else ''
            if not name or name.lower() in seen_names:
                continue
            seen_names.add(name.lower())

            grade = texts[grade_col].strip() if grade_col >= 0 and grade_col < len(texts) else ''
            specialty = texts[specialty_col].strip() if specialty_col >= 0 and specialty_col < len(texts) else ''
            ecue_raw = texts[ecue_col].strip() if ecue_col >= 0 and ecue_col < len(texts) else ''
            establishment = texts[establishment_col].strip() if establishment_col >= 0 and establishment_col < len(texts) else ''

            # Parse ECUE list (comma-separated)
            ecues = [e.strip() for e in ecue_raw.split(',') if e.strip()]

            teachers.append({
                'name': name,
                'grade': grade,
                'specialty': specialty,
                'establishment': establishment,
                'ecues': ecues,
            })

    logger.info("Extracted %d teachers", len(teachers))
    return teachers


def save_extracted_data(program_id: int, extracted: dict) -> dict:
    """
    Persist extracted AAP, Competences, and their matrix links to the database.

    Clears any existing AAP / competence data for the program before inserting,
    so the operation is idempotent.

    Returns:
        {'aaps_count': N, 'competences_count': N, 'links_count': N}
    """
    from app import db
    from app.models.program_learning import (
        ProgramAAP,
        ProgramCompetence,
        aap_competence_link,
    )

    logger.info("Saving extracted data for program_id=%d", program_id)

    # ── 1. Delete existing data ──────────────────────────────────────────
    # Must delete M2M link rows FIRST (bulk delete doesn't trigger cascade)
    existing_aap_ids = [a.id for a in ProgramAAP.query.filter_by(program_id=program_id).all()]
    existing_comp_ids = [c.id for c in ProgramCompetence.query.filter_by(program_id=program_id).all()]
    if existing_aap_ids or existing_comp_ids:
        db.session.execute(
            aap_competence_link.delete().where(
                aap_competence_link.c.aap_id.in_(existing_aap_ids) |
                aap_competence_link.c.competence_id.in_(existing_comp_ids)
            )
        )
    ProgramAAP.query.filter_by(program_id=program_id).delete()
    ProgramCompetence.query.filter_by(program_id=program_id).delete()
    db.session.flush()

    # ── 2. Insert competences ────────────────────────────────────────────
    comp_map: Dict[str, ProgramCompetence] = {}
    for item in extracted.get('competences', []):
        comp = ProgramCompetence(
            program_id=program_id,
            code=item['code'],
            description=item.get('description', ''),
        )
        db.session.add(comp)
        comp_map[item['code']] = comp

    db.session.flush()  # assign ids

    # ── 3. Insert AAPs ──────────────────────────────────────────────────
    aap_map: Dict[str, ProgramAAP] = {}
    for item in extracted.get('aaps', []):
        desc = item.get('description') or item.get('denomination', '')
        aap = ProgramAAP(
            program_id=program_id,
            code=item['code'],
            description=desc,
            order=item.get('order', 0),
        )
        db.session.add(aap)
        aap_map[item['code']] = aap

    db.session.flush()  # assign ids

    # ── 4. Create matrix links ──────────────────────────────────────────
    links_count = 0
    for entry in extracted.get('matrix', []):
        comp_code = entry['competence_code']
        comp = comp_map.get(comp_code)
        if comp is None:
            logger.warning("Matrix references unknown competence %s — skipped", comp_code)
            continue
        for aap_code in entry.get('aap_codes', []):
            aap = aap_map.get(aap_code)
            if aap is None:
                logger.warning("Matrix references unknown AAP %s — skipped", aap_code)
                continue
            aap.competences.append(comp)
            links_count += 1

    db.session.commit()

    # ── 5. Auto-link courses matching ECUE names from study plan ────────
    courses_linked = 0
    study_plan = extracted.get('study_plan', [])
    if study_plan:
        from app.models.courses import Course
        from app.models.institutions import Program, program_course
        program = Program.query.get(program_id)
        if program:
            existing_course_ids = {c.id for c in program.courses}
            for ecue in study_plan:
                ecue_name = ecue.get('name', '').strip()
                if not ecue_name:
                    continue
                # Match by exact title or case-insensitive containment
                course = Course.query.filter(
                    db.func.lower(Course.title) == ecue_name.lower()
                ).first()
                if course and course.id not in existing_course_ids:
                    program.courses.append(course)
                    existing_course_ids.add(course.id)
                    courses_linked += 1
                    logger.info("Auto-linked course '%s' (id=%d) to program %d", course.title, course.id, program_id)
            db.session.commit()

    result = {
        'aaps_count': len(aap_map),
        'competences_count': len(comp_map),
        'links_count': links_count,
        'courses_linked': courses_linked,
    }
    logger.info("Save complete: %s", result)
    return result


# ═══════════════════════════════════════════════════════════════════════════════
# ORCHESTRATOR PIPELINE — process a full formation descriptor
# ═══════════════════════════════════════════════════════════════════════════════

def _normalize(text: str) -> str:
    """Lowercase + collapse whitespace for fuzzy matching."""
    return ' '.join(text.lower().split())


def process_program_descriptor(program_id: int, file_path: str) -> dict:
    """
    Full agentic pipeline: extract everything from the .docx descriptor,
    create courses, teachers, link them, and save to DB.

    Returns a detailed report of all operations performed.
    """
    from app import db
    from app.models.courses import Course
    from app.models.institutions import Program, program_course
    from app.models.users import User
    from app.models.program_learning import ProgramAAP, ProgramCompetence, aap_competence_link
    from werkzeug.security import generate_password_hash
    import secrets

    steps = []  # each step: {'agent': str, 'status': str, 'details': dict}

    # ── STEP 1: Extract everything from the document ──────────────────────
    logger.info("[Pipeline] Starting extraction for program %d", program_id)
    extracted = extract_program_descriptor(file_path, program_id)
    steps.append({
        'agent': 'Extraction du document',
        'status': 'success',
        'details': {
            'aaps': len(extracted.get('aaps', [])),
            'competences': len(extracted.get('competences', [])),
            'matrix': len(extracted.get('matrix', [])),
            'modules': len(extracted.get('study_plan', [])),
            'teachers': len(extracted.get('teachers', [])),
        }
    })

    # ── STEP 2: Save AAP, Competences, Matrix ─────────────────────────────
    logger.info("[Pipeline] Saving AAP/Competences/Matrix")
    save_result = save_extracted_data(program_id, extracted)
    steps.append({
        'agent': 'Enregistrement AAP & Compétences',
        'status': 'success',
        'details': save_result,
    })

    # ── STEP 3: Create teachers ──────────────────────────────────────────
    logger.info("[Pipeline] Creating teachers")
    teachers_data = extracted.get('teachers', [])
    teacher_map: Dict[str, User] = {}  # normalized ECUE name → teacher user
    teachers_created = []
    teachers_existing = []

    for t in teachers_data:
        full_name = t['name'].strip()
        if not full_name:
            continue

        # Generate username from name: "Salhi Jamil" → "salhi.jamil"
        parts = full_name.split()
        username = '.'.join(p.lower() for p in parts)

        # Check if teacher already exists
        user = User.query.filter(
            db.func.lower(User.username) == username.lower()
        ).first()

        if not user:
            password = secrets.token_urlsafe(8)
            user = User(
                username=username,
                email=f"{username}@esb-learning.tn",
                password_hash=generate_password_hash(password),
                is_teacher=True,
                is_superuser=False,
            )
            db.session.add(user)
            db.session.flush()
            teachers_created.append({
                'name': full_name,
                'username': username,
                'password': password,
                'email': user.email,
                'id': user.id,
            })
            logger.info("[Pipeline] Created teacher: %s (id=%d)", username, user.id)
        else:
            teachers_existing.append({
                'name': full_name,
                'username': user.username,
                'id': user.id,
            })

        # Map each ECUE to this teacher
        for ecue_name in t.get('ecues', []):
            norm = _normalize(ecue_name)
            if norm:
                teacher_map[norm] = user

    db.session.flush()
    steps.append({
        'agent': 'Création des enseignants',
        'status': 'success',
        'details': {
            'created': len(teachers_created),
            'existing': len(teachers_existing),
            'teachers_created': teachers_created,
            'teachers_existing': teachers_existing,
        },
    })

    # ── STEP 4: Create courses for each ECUE ─────────────────────────────
    logger.info("[Pipeline] Creating courses from study plan")
    study_plan = extracted.get('study_plan', [])
    program = db.session.get(Program, program_id)
    existing_course_titles = {_normalize(c.title) for c in program.courses} if program else set()

    # We need a fallback teacher for courses without a mapped teacher
    admin_user = User.query.filter_by(is_superuser=True).first()
    fallback_teacher_id = admin_user.id if admin_user else 1

    courses_created = []
    courses_existing = []
    modules_table = []  # comprehensive table data

    for ecue in study_plan:
        ecue_name = ecue.get('name', '').strip()
        ecue_code = ecue.get('code', '').strip()
        semester = ecue.get('semester', 0)
        ue = ecue.get('ue', '')
        if not ecue_name:
            continue

        norm_name = _normalize(ecue_name)

        # Find or create course
        course = Course.query.filter(
            db.func.lower(Course.title) == ecue_name.lower()
        ).first()

        # Find the teacher for this ECUE
        assigned_teacher = teacher_map.get(norm_name)
        # Also try partial matching if exact doesn't work
        if not assigned_teacher:
            for t_ecue, t_user in teacher_map.items():
                if t_ecue in norm_name or norm_name in t_ecue:
                    assigned_teacher = t_user
                    break

        teacher_id = assigned_teacher.id if assigned_teacher else fallback_teacher_id

        if not course:
            course = Course(
                title=ecue_name,
                description=f"Semestre {semester} — {ue}" if ue else f"Semestre {semester}",
                teacher_id=teacher_id,
            )
            db.session.add(course)
            db.session.flush()

            # Link to program
            if program and norm_name not in existing_course_titles:
                program.courses.append(course)
                existing_course_titles.add(norm_name)

            courses_created.append({
                'id': course.id,
                'title': ecue_name,
                'code': ecue_code,
                'semester': semester,
                'ue': ue,
                'teacher': assigned_teacher.username if assigned_teacher else '—',
            })
            logger.info("[Pipeline] Created course: %s (id=%d, teacher=%s)",
                        ecue_name, course.id, assigned_teacher.username if assigned_teacher else 'admin')
        else:
            # Link existing course to program if not already linked
            if program and norm_name not in existing_course_titles:
                program.courses.append(course)
                existing_course_titles.add(norm_name)

            # Update teacher if we found one and course has admin/fallback teacher
            if assigned_teacher and course.teacher_id == fallback_teacher_id:
                course.teacher_id = teacher_id

            courses_existing.append({
                'id': course.id,
                'title': course.title,
                'code': ecue_code,
                'semester': semester,
            })

        # Build comprehensive table row
        modules_table.append({
            'course_id': course.id,
            'title': ecue_name,
            'code': ecue_code,
            'semester': semester,
            'ue': ue,
            'teacher_id': course.teacher_id,
            'teacher_name': assigned_teacher.username if assigned_teacher else (
                course.teacher.username if course.teacher else '—'
            ),
            'course_link': f"/courses/{course.id}",
        })

    db.session.commit()

    steps.append({
        'agent': 'Création des modules/cours',
        'status': 'success',
        'details': {
            'created': len(courses_created),
            'existing': len(courses_existing),
            'courses_created': courses_created,
        },
    })

    # ── STEP 5: Summary ──────────────────────────────────────────────────
    steps.append({
        'agent': 'Finalisation',
        'status': 'success',
        'details': {
            'total_modules': len(modules_table),
            'total_teachers': len(teachers_created) + len(teachers_existing),
        },
    })

    logger.info("[Pipeline] Pipeline complete for program %d", program_id)

    # ── STEP 6: Extract syllabi for courses with documents ──────────────
    logger.info("[Pipeline] Attempting syllabus extraction for linked courses")
    syllabus_results = []
    try:
        from app.models.syllabus import Syllabus
        from app.models.documents import Document

        for module in modules_table:
            course_id = module.get('course_id')
            if not course_id:
                continue

            existing_syllabus = Syllabus.query.filter_by(course_id=course_id).first()

            docs = Document.query.filter_by(course_id=course_id).all()
            pdf_docs = [d for d in docs if d.file_path and d.file_path.lower().endswith('.pdf')]

            syllabus_results.append({
                'course_id': course_id,
                'title': module.get('title', ''),
                'has_syllabus': existing_syllabus is not None,
                'pdf_count': len(pdf_docs),
            })
    except Exception as e:
        logger.warning("[Pipeline] Syllabus scan failed: %s", e)

    steps.append({
        'agent': 'Analyse des syllabus',
        'status': 'success',
        'details': {
            'courses_scanned': len(syllabus_results),
            'with_syllabus': len([r for r in syllabus_results if r['has_syllabus']]),
            'with_pdfs': len([r for r in syllabus_results if r['pdf_count'] > 0]),
        },
    })

    return {
        'steps': steps,
        'modules_table': modules_table,
        'teachers_created': teachers_created,
    }
