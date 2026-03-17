import os
import json
import pdfplumber
import markdownify
import camelot
import re  # For section extraction and CLO parsing
from datetime import datetime
from app import db
from app.models import Syllabus, Course

from flask import current_app
import logging
import time  # For retries
from app.models import Document, Course  # <-- Add this import




# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)


from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage, SystemMessage

class SyllabusService:
    @staticmethod
    def _get_client():
        # Returns correctly configured Gemini LangChain client
        api_key = current_app.config.get("GOOGLE_API_KEY")
        model = current_app.config.get("GEMINI_MODEL")
        if not api_key:
            logger.error("Google API key is not configured")
            raise ValueError("Google API key is not configured")
        
        return ChatGoogleGenerativeAI(
            model=model,
            google_api_key=api_key
        )


    @staticmethod
    def estimate_tokens(text):
        """Rough token estimate: ~4 chars per token + prompt overhead."""
        return int(len(text) / 4) + 500  # Conservative


    @staticmethod
    def extract_section(text, start_keywords, end_keywords=None, max_chars=8000):
        """Extract a section from text using keywords."""
        text_lower = text.lower()
        start_idx = -1
        for kw in start_keywords:
            idx = text_lower.find(kw.lower())
            if idx != -1:
                start_idx = max(start_idx, idx) if start_idx != -1 else idx
        if start_idx == -1:
            return text[:max_chars]  # Fallback to prefix


        end_idx = len(text)
        if end_keywords:
            for kw in end_keywords:
                idx = text_lower.find(kw.lower(), start_idx)
                if idx != -1:
                    end_idx = min(end_idx, idx)
        section = text[start_idx:end_idx][:max_chars]
        logger.debug(f"Extracted section ({len(start_keywords)} keywords): {len(section)} chars")
        return section


    # ---------------- CRUD & Persistence ----------------
    @staticmethod
    def get_syllabus_by_course(course_id):
        """
        Retrieve the syllabus for a specific course.
        ✅ FIXED: This properly queries the database every time.
        """
        try:
            # Query directly from database - don't cache
            syllabus = Syllabus.query.filter_by(course_id=course_id).first()
            
            if syllabus:
                logger.debug(f"✅ Syllabus retrieved for course {course_id}")
                logger.debug(f"   - file_path: {syllabus.file_path}")
                logger.debug(f"   - has_clo_data: {bool(syllabus.clo_data)}")
                logger.debug(f"   - has_weekly_plan: {bool(syllabus.weekly_plan)}")
                return syllabus
            else:
                logger.debug(f"⚠️ No syllabus found for course {course_id}")
                return None
                
        except Exception as e:
            logger.error(f"❌ Error retrieving syllabus for course {course_id}: {str(e)}")
            return None


    @staticmethod
    def create_syllabus(course_id, syllabus_type='bga', clo_data=None, plo_data=None, weekly_plan=None, clo_stats=None, file_path=None, tn_data=None):
        syllabus = Syllabus(
            course_id=course_id,
            syllabus_type=syllabus_type,
            clo_data=clo_data or [],
            plo_data=plo_data or [],
            weekly_plan=weekly_plan or [],
            clo_stats=clo_stats or {},
            file_path=file_path,
            tn_data=tn_data
        )
        db.session.add(syllabus)
        try:
            db.session.commit()
            logger.debug(f"Created {syllabus_type} syllabus for course {course_id}: file_path='{file_path}'")
        except Exception as e:
            db.session.rollback()
            logger.error(f"Failed to create {syllabus_type} syllabus for course {course_id}: {str(e)}")
            raise
        return syllabus


    @staticmethod
    def update_syllabus(course_id, clo_data=None, plo_data=None, weekly_plan=None, clo_stats=None, file_path=None, tn_data=None):
        syllabus = SyllabusService.get_syllabus_by_course(course_id)
        if not syllabus:
            logger.warning(f"No syllabus to update for course {course_id}")
            return None
        if clo_data is not None:
            syllabus.clo_data = clo_data
        if plo_data is not None:
            syllabus.plo_data = plo_data
        if weekly_plan is not None:
            syllabus.weekly_plan = weekly_plan
        if clo_stats is not None:
            syllabus.clo_stats = clo_stats
        if file_path is not None:
            syllabus.file_path = file_path
        if tn_data is not None:
            syllabus.tn_data = tn_data
        syllabus.updated_at = datetime.utcnow()
        try:
            db.session.commit()
            logger.debug(f"Updated syllabus for course {course_id}: file_path='{syllabus.file_path}', weekly_plan length={len(weekly_plan) if weekly_plan else 0}")
        except Exception as e:
            db.session.rollback()
            logger.error(f"Failed to update syllabus for course {course_id}: {str(e)}")
            raise
        return syllabus


    @staticmethod
    def get_weekly_plan(course_id):
        """Load weekly_plan from DB (for persistence in view_course)."""
        syllabus = SyllabusService.get_syllabus_by_course(course_id)
        if syllabus and syllabus.weekly_plan:
            return syllabus.weekly_plan  # Already JSON list of dicts
        return []


    @staticmethod
    def update_week_data(course_id, week_num, updated_data):
        """Update specific week in weekly_plan (for edit_week route)."""
        syllabus = SyllabusService.get_syllabus_by_course(course_id)
        if not syllabus or not syllabus.weekly_plan:
            logger.warning(f"No weekly_plan to update for course {course_id}")
            return False
        weekly_plan = syllabus.weekly_plan  # List of dicts
        if 0 <= week_num - 1 < len(weekly_plan):
            weekly_plan[week_num - 1].update(updated_data)
            syllabus.weekly_plan = weekly_plan
            syllabus.updated_at = datetime.utcnow()
            try:
                db.session.commit()
                logger.debug(f"Updated week {week_num} for course {course_id}")
                return True
            except Exception as e:
                db.session.rollback()
                logger.error(f"Failed to update week {week_num}: {str(e)}")
        return False


    @staticmethod
    def get_week_data(course_id, week_num):
        """Fetch specific week data from DB."""
        syllabus = SyllabusService.get_syllabus_by_course(course_id)
        if not syllabus or not syllabus.weekly_plan:
            logger.warning(f"No weekly_plan for course {course_id}")
            return None
        week_data = next((w for w in syllabus.weekly_plan if w.get("Week#") == week_num), None)
        if not week_data:
            logger.warning(f"Week {week_num} not found for course {course_id}")
        return week_data


    @staticmethod
    def calculate_week_clo_percentages(course_id, week_num):
        """Calculate CLO percentages for a specific week (equal distribution)."""
        week_data = SyllabusService.get_week_data(course_id, week_num)
        if not week_data or not week_data.get("Related CLOs"):
            logger.warning(f"No CLOs for week {week_num} in course {course_id}")
            return {}
        clos = week_data["Related CLOs"]
        total_clos = len(clos)
        if total_clos == 0:
            return {}
       
        base_pct = 100 / total_clos
        percentages = {f"CLO{clo}": round(base_pct, 2) for clo in clos}
        logger.debug(f"Week {week_num} CLO percentages for course {course_id}: {percentages}")
        return percentages
   
    @staticmethod
    def calculate_clo_coverage_stats(course_id):
        """
        Calculate CLO coverage % based on weekly_plan (persists to DB).
        """
        syllabus = SyllabusService.get_syllabus_by_course(course_id)
        if not syllabus or not syllabus.weekly_plan:
            logger.warning(f"No weekly_plan for CLO stats on course {course_id}")
            return {}


        clo_counts = {}
        total_mentions = 0


        for week_data in syllabus.weekly_plan:
            related_clos = week_data.get("Related CLOs", [])
            if not related_clos:
                continue
           
            # Handle list or string
            if isinstance(related_clos, str):
                nums = [int(x) for x in re.findall(r"\d+", related_clos)]
            else:
                nums = [int(clo) for clo in related_clos if isinstance(clo, (int, float))]


            for clo_num in nums:
                clo_key = f"CLO{clo_num}"
                clo_counts[clo_key] = clo_counts.get(clo_key, 0) + 1
                total_mentions += 1


        if total_mentions == 0:
            return {}


        # Calculate and round percentages (sum to 100)
        raw_percentages = {}
        for clo_key, weeks in clo_counts.items():
            raw_pct = (weeks / total_mentions) * 100
            raw_percentages[clo_key] = {"weeks": weeks, "raw_pct": raw_pct}


        floored = {k: {"weeks": v["weeks"], "percentage": int(v["raw_pct"] // 1)} for k, v in raw_percentages.items()}
        total_floored = sum(v["percentage"] for v in floored.values())
        remainder = 100 - total_floored


        if remainder != 0:
            fractions = [(k, v["raw_pct"] - int(v["raw_pct"] // 1)) for k, v in raw_percentages.items()]
            fractions.sort(key=lambda x: x[1], reverse=True)
            for i in range(abs(remainder)):
                clo_key = fractions[i][0]
                if remainder > 0:
                    floored[clo_key]["percentage"] += 1
                else:
                    floored[clo_key]["percentage"] = max(0, floored[clo_key]["percentage"] - 1)


        clo_stats = {k: {"weeks": v["weeks"], "percentage": v["percentage"]} for k, v in floored.items()}


        # Save to DB
        syllabus.clo_stats = clo_stats
        try:
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            logger.error(f"Failed to save CLO stats: {str(e)}")
            raise


        return clo_stats


    @staticmethod
    def _extract_text_from_docx(file_path):
        """Extract text and tables from a DOCX file, returning markdown-style text."""
        import docx as python_docx
        doc = python_docx.Document(file_path)
        parts = []

        for element in doc.element.body:
            tag = element.tag.split('}')[-1]
            if tag == 'p':
                # Paragraph
                para_text = element.text_content() if hasattr(element, 'text_content') else ''
                # Use python-docx paragraph
            elif tag == 'tbl':
                pass  # handled separately

        # Paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style = para.style.name if para.style else ''
                if 'Heading 1' in style:
                    parts.append(f"# {text}")
                elif 'Heading 2' in style:
                    parts.append(f"## {text}")
                elif 'Heading 3' in style:
                    parts.append(f"### {text}")
                else:
                    parts.append(text)

        full_text = "\n\n".join(parts)

        # Tables
        table_parts = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                table_parts.append("\n".join(rows))

        full_tables = "\n\n".join(table_parts)
        return full_text, full_tables

    @staticmethod
    def extract_from_file(pdf_path):
        if not os.path.exists(pdf_path):
            logger.error(f"File not found at {pdf_path}")
            raise FileNotFoundError(f"File not found at {pdf_path}")

        ext = os.path.splitext(pdf_path)[1].lower()
        markdown_output = []
        full_tables = ""

        try:
            if ext in ('.docx', '.doc'):
                # DOCX extraction
                logger.debug(f"Extracting from DOCX file: {pdf_path}")
                full_markdown, full_tables = SyllabusService._extract_text_from_docx(pdf_path)
                logger.debug(f"DOCX text extracted (length: {len(full_markdown)} chars)")
            else:
                # PDF extraction with pdfplumber
                logger.debug(f"Extracting from PDF file: {pdf_path}")
                with pdfplumber.open(pdf_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            md_text = markdownify.markdownify(text, heading_style="ATX")
                            markdown_output.append(md_text)

                full_markdown = "\n\n".join(markdown_output)
                logger.debug(f"PDF text extracted to markdown (length: {len(full_markdown)} chars)")

                # Extract tables with camelot
                try:
                    tables = camelot.read_pdf(pdf_path, pages='all')
                    table_data = []
                    for table in tables:
                        table_data.append(table.df.to_string())
                    full_tables = "\n\n".join(table_data)
                    logger.debug(f"Extracted {len(tables)} tables from PDF")
                except AttributeError as e:
                    logger.warning(f"Failed to extract tables with camelot: {str(e)}. Proceeding with text only.")
                    full_tables = ""
                except Exception as e:
                    logger.warning(f"Unexpected error during table extraction: {str(e)}. Proceeding with text only.")
                    full_tables = ""

            # Combine text and tables for LLM input
            combined_text = full_markdown + ("\n\nExtracted Tables:\n" + full_tables if full_tables else "")
            logger.debug(f"Full combined text length for LLM: {len(combined_text)} chars (~{SyllabusService.estimate_tokens(combined_text)} tokens)")
        except Exception as e:
            logger.error(f"Failed to extract content from file {pdf_path}: {str(e)}")
            raise


        # Create LLM client ONCE and reuse for all calls
        llm = SyllabusService._get_client()

        def ask_llm(prompt_template, text, section="general", max_retries=2):
            for attempt in range(max_retries + 1):
                try:
                    estimated_tokens = SyllabusService.estimate_tokens(text + prompt_template)
                    if estimated_tokens > 10000:  # Pre-check to avoid 413
                        logger.warning(f"{section} input too large ({estimated_tokens} tokens), reducing...")
                        text = text[:6000]  # Truncate if still over
                        estimated_tokens = SyllabusService.estimate_tokens(text + prompt_template)



                    # Format the prompt with the text
                    try:
                        user_content = prompt_template.format(combined_text=text)
                    except KeyError:
                        # Fallback if template key doesn't match
                        user_content = f"{prompt_template}\n\n{text}"

                    messages = [
                        SystemMessage(content=(
                            "Tu es un assistant d'extraction de syllabus. "
                            "Ton rôle est d'extraire le texte EXACTEMENT tel qu'il apparaît dans le document, "
                            "sans paraphraser, interpréter ou modifier le contenu. "
                            "Retourne UNIQUEMENT un JSON valide, sans texte avant ou après. "
                            "Si aucune donnée ne peut être extraite, retourne une liste vide []. "
                            "Assure-toi d'extraire TOUS les éléments disponibles dans le texte fourni."
                        )),
                        HumanMessage(content=user_content)
                    ]
                    
                    logger.debug(f"Sending {section} prompt with {len(user_content)} chars (~{estimated_tokens} tokens)")
                    response = llm.invoke(messages)
                    content = response.content.strip()
                    logger.debug(f"Raw LLM response for {section}: {content[:1000]}...")
                    try:
                        # Robust JSON extraction: handle markdown code blocks
                        json_str = content
                        
                        # Try stripping ```json ... ``` or ``` ... ```
                        import re
                        code_block_match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", content)
                        if code_block_match:
                            json_str = code_block_match.group(1).strip()
                        elif content.startswith("{") or content.startswith("["):
                            # Raw JSON
                            json_str = content
                        else:
                            # Try to find outermost braces
                            s_brace = content.find("{")
                            e_brace = content.rfind("}")
                            s_bracket = content.find("[")
                            e_bracket = content.rfind("]")
                            if s_bracket != -1 and e_bracket > s_bracket and (s_brace == -1 or s_bracket < s_brace):
                                json_str = content[s_bracket:e_bracket+1]
                            elif s_brace != -1 and e_brace > s_brace:
                                json_str = content[s_brace:e_brace+1]
                        
                        json_content = json.loads(json_str)
                        logger.debug(f"Parsed {section} response as JSON: {json_content[:2] if isinstance(json_content, list) else json_content}")
                        return json_content
                    except json.JSONDecodeError as e:
                        logger.error(f"LLM {section} response is not valid JSON: {content[:500]}... Error: {str(e)}")
                        return []
                except Exception as e:
                    if "413" in str(e) or "rate_limit_exceeded" in str(e):
                        wait_time = 10 * (attempt + 1)  # 10s, 20s max (was 30s, 60s)
                        logger.warning(f"{section} failed (413/429), retrying in {wait_time}s (attempt {attempt + 1}/{max_retries + 1})")
                        time.sleep(wait_time)
                        if attempt == max_retries:
                            logger.error(f"Max retries exceeded for {section}: {str(e)}")
                            return []
                    else:
                        logger.error(f"LLM {section} request failed: {str(e)}")
                        return []
            return []


        # Enhanced prompts (same as before)
        plo_prompt = (
            "Tu es un assistant d'extraction de syllabus.\n"
            "Analyse le texte suivant et extrait TOUS les PLOs listés (de PLO# 1 à la fin).\n"
            "Retourne UNIQUEMENT un JSON valide, sans texte avant/après.\n\n"
            "Le JSON doit être une liste d'objets avec les clés :\n"
            "- \"PLO#\" : entier\n"
            "- \"PLO Description\" : texte exact (sans modification)\n\n"
            "Texte :\n"
            "{combined_text}"
        )


        clo_prompt = (
            "Tu es un assistant d'extraction de syllabus.\n"
            "Analyse le texte suivant et extrait TOUS les CLOs de la section 'Course Learning Outcomes (CLOs)'.\n"
            "Pour chaque CLO, parse les PLOs liés (ex: '2,4,5' -> [2,4,5]). Si table messy, utilise le contexte.\n"
            "Retourne UNIQUEMENT un JSON valide, sans texte avant/après.\n\n"
            "Le JSON doit être une liste d'objets avec les clés :\n"
            "- \"CLO#\" : entier\n"
            "- \"CLO Description\" : texte exact\n"
            "- \"Linked PLOs\" : liste d'entiers (tous les PLOs liés, ex: [2,4,5])\n\n"
            "Texte :\n"
            "{combined_text}"
        )


        weekly_prompt = (
            "Tu es un assistant d'extraction de syllabus.\n"
            "Retourne UNIQUEMENT un JSON valide, sans texte avant ni après.\n"
            "Si tu ne peux pas extraire, retourne [].\n\n"
            "Extrait TOUS les éléments du 'Course weekly plan' (de Week# 1 à 14).\n"
            "Parse la table complète: Week#, Topic, Class Objectives, Related CLOs (ex: '1,2' -> [1,2]), Activities/Assessment, Assignments/Readings.\n"
            "Si texte tronqué ou messy, extrait ce qui est disponible et liste TOUS les weeks visibles.\n\n"
            "Le JSON doit être une liste d'objets avec les clés :\n"
            "- \"Week#\" : entier\n"
            "- \"Topic\" : texte exact\n"
            "- \"Class Objectives\" : texte exact\n"
            "- \"Related CLOs\" : liste d'entiers (ex: [1,2])\n"
            "- \"Activities/Assessment\" : texte exact\n"
            "- \"Assignments/Readings\" : texte exact\n\n"
            "Document :\n"
            "{combined_text}"
        )


        clo_data = plo_data = weekly_plan = []

        # Run PLO and CLO extractions IN PARALLEL to halve wait time
        import concurrent.futures

        plo_section = SyllabusService.extract_section(combined_text, ["plo#", "program learning outcomes"], end_keywords=["clo", "course learning outcomes"])
        clo_section = SyllabusService.extract_section(combined_text, ["course learning outcomes (clos)", "clo#"], end_keywords=["weekly plan", "pedagogical"])

        def extract_plo():
            try:
                result = ask_llm(plo_prompt, plo_section, "PLO")
                return result if isinstance(result, list) else []
            except Exception as e:
                logger.error(f"Error processing PLO data: {str(e)}")
                return []

        def extract_clo():
            try:
                result = ask_llm(clo_prompt, clo_section, "CLO")
                if not result:
                    logger.warning("CLO extraction empty - retrying with full text chunk")
                    result = ask_llm(clo_prompt, combined_text[:8000], "CLO-fallback")
                return result if isinstance(result, list) else []
            except Exception as e:
                logger.error(f"Error processing CLO data: {str(e)}")
                return []

        with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
            future_plo = executor.submit(extract_plo)
            future_clo = executor.submit(extract_clo)
            plo_data = future_plo.result()
            clo_data = future_clo.result()

        logger.debug(f"Parallel extraction done — PLO: {len(plo_data)}, CLO: {len(clo_data)}")

        # Extract Weekly Plan section (after PLO+CLO complete)
        weekly_section = SyllabusService.extract_section(combined_text, ["course weekly plan", "week #"], end_keywords=["educational resources", "classroom policy", "reference books"])
        try:
            weekly_response = ask_llm(weekly_prompt, weekly_section, "weekly")
            weekly_plan = weekly_response if isinstance(weekly_response, list) else []
            if len(weekly_plan) < 2:
                logger.warning(f"Weekly plan incomplete ({len(weekly_plan)} items) - retrying with larger chunk")
                larger_weekly = combined_text[combined_text.lower().find("course weekly plan") - 2000:]  # Include some context before
                weekly_response = ask_llm(weekly_prompt, larger_weekly[:10000], "weekly-fallback")
                weekly_plan = weekly_response if isinstance(weekly_response, list) else weekly_plan
            logger.debug(f"Parsed weekly plan: {len(weekly_plan)} items")
        except Exception as e:
            logger.error(f"Error processing weekly plan: {str(e)}")
            weekly_plan = []


        return {
            "clo_data": clo_data,
            "plo_data": plo_data,
            "weekly_plan": weekly_plan
        }


    @staticmethod

    def extract_syllabus(course_id):
        """On-demand extraction (if weekly_plan missing; updates DB)."""
        syllabus = SyllabusService.get_syllabus_by_course(course_id)
        if not syllabus or not syllabus.file_path:
            logger.warning(f"No file for extraction on course {course_id}")
            return {"error": "No syllabus file available"}


        full_path = os.path.join(current_app.config['UPLOAD_FOLDER'], syllabus.file_path)
        try:
            extracted = SyllabusService.extract_from_file(full_path)
            # Update DB with extracted data
            SyllabusService.update_syllabus(
                course_id=course_id,
                clo_data=extracted["clo_data"],
                plo_data=extracted["plo_data"],
                weekly_plan=extracted["weekly_plan"]
            )
            # Calculate stats post-extraction
            SyllabusService.calculate_clo_coverage_stats(course_id)
            logger.debug(f"On-demand extraction successful for course {course_id}")
            return {"success": True, "weekly_plan": extracted["weekly_plan"]}
        except Exception as e:
            logger.error(f"Extraction failed for course {course_id}: {str(e)}")
            return {"error": str(e)}


    @staticmethod
    def delete_week_attachment(attachment_id, user_id):
        """
        Delete a week attachment if the user is the course teacher.
        Returns True on success, False otherwise.
        """
        attachment = Document.query.filter_by(id=attachment_id).first()
        if not attachment:
            logger.warning(f"Attachment {attachment_id} not found")
            return False


        course = Course.query.get(attachment.course_id)
        if not course or course.teacher_id != user_id:
            logger.warning(f"User   {user_id} not authorized to delete attachment {attachment_id} for course {course.id if course else 'N/A'}")
            return False


        # Delete physical file if exists and not a quiz
        if attachment.file_path and not attachment.is_quiz:
            full_file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], attachment.file_path)
            if os.path.exists(full_file_path):
                try:
                    os.remove(full_file_path)
                    logger.debug(f"Deleted physical file: {full_file_path}")
                except OSError as e:
                    logger.error(f"Failed to delete physical file {full_file_path}: {str(e)}")


        try:
            db.session.delete(attachment)
            db.session.commit()
            logger.debug(f"Deleted attachment {attachment_id} from DB for course {course.id}, week {attachment.week_number}")
            return True
        except Exception as e:
            db.session.rollback()
            logger.error(f"Database error deleting attachment {attachment_id}: {str(e)}")
            return False


